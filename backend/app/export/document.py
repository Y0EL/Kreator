from __future__ import annotations

import re
from io import BytesIO

_LONG_TOKEN = re.compile(r"\S{56,}")
_FONT = "Times New Roman"
_ACCENT = "FF2F2F"


def _break_long_tokens(text: str, n: int = 55) -> str:
    def rep(m: re.Match[str]) -> str:
        s = m.group(0)
        return " ".join(s[i : i + n] for i in range(0, len(s), n))

    return _LONG_TOKEN.sub(rep, text)


def _is_stage(line: str) -> bool:
    return line.startswith("[") and line.endswith("]") and not line.startswith("[SEGMEN")


def _bottom_rule(paragraph, sz: str = "18") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), _ACCENT)
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx(title: str, body: str, sources: str, meta_lines: list[str]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    black = RGBColor(0, 0, 0)
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = black

    def para(text, size=12, bold=False, italic=False, after=6, justify=False, upper=False):
        p = doc.add_paragraph()
        r = p.add_run(text.upper() if upper else text)
        r.font.name = _FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = black
        p.paragraph_format.space_after = Pt(after)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    _bottom_rule(para(title, size=22, bold=True, after=2))
    for m in meta_lines:
        para(m, size=10, italic=True, after=1)
    para("", after=4)

    for raw in body.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("[SEGMEN"):
            _bottom_rule(para(line.strip("[]"), size=14, bold=True, after=4, upper=True), sz="6")
        elif _is_stage(line):
            para(line, italic=True)
        else:
            para(line, after=8, justify=True)

    if sources:
        doc.add_page_break()
        _bottom_rule(para("SUMBER DAN BUKTI", size=15, bold=True, after=4), sz="6")
        for raw in sources.split("\n"):
            if raw.strip():
                para(raw, size=10, after=2)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_pdf(title: str, body: str, sources: str, meta_lines: list[str]) -> bytes:
    from fpdf.enums import XPos, YPos
    from fpdf.fpdf import FPDF

    def enc(text: str) -> str:
        return _break_long_tokens(text).encode("latin-1", "replace").decode("latin-1")

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(True, margin=16)
    pdf.set_margins(18, 16, 18)
    pdf.add_page()
    pdf.set_text_color(0, 0, 0)

    def rule(thick: float = 0.7) -> None:
        pdf.set_draw_color(255, 47, 47)
        pdf.set_line_width(thick)
        y = pdf.get_y() + 1
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(3.5)

    def cell(text: str, size: int, style: str = "", h: float = 6.0) -> None:
        pdf.set_font("Times", style, size)
        pdf.set_text_color(0, 0, 0)
        pdf.multi_cell(0, h, enc(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    cell(title, 20, "B", h=8)
    rule(1.0)
    for m in meta_lines:
        cell(m, 10, "I", h=5)
    pdf.ln(2)

    for raw in body.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("[SEGMEN"):
            pdf.ln(2)
            cell(line.strip("[]").upper(), 14, "B", h=7)
            rule(0.5)
        elif _is_stage(line):
            cell(line, 11, "I")
        else:
            cell(line, 12, "", h=6.4)

    if sources:
        pdf.add_page()
        cell("SUMBER DAN BUKTI", 15, "B", h=7)
        rule(0.5)
        cell(sources, 10, "", h=5)

    return bytes(pdf.output())  # type: ignore[arg-type]
